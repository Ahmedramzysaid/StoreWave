from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def shd(cell, h):
    tc = cell._tc
    p = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), h)
    p.append(s)


def ah(doc, txt, lv=1, col=None):
    hd = doc.add_heading(txt, level=lv)
    hd.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if hd.runs and col:
        hd.runs[0].font.color.rgb = RGBColor(*bytes.fromhex(col))
    return hd


def ap(doc, txt, bold=False, italic=False, sz=11):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(sz)
    return p


def at(doc, hdrs, rows, hc='1F4E79'):
    t = doc.add_table(rows=1 + len(rows), cols=len(hdrs))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hc_cells = t.rows[0].cells
    for i, h in enumerate(hdrs):
        hc_cells[i].text = h
        shd(hc_cells[i], hc)
        r = hc_cells[i].paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    for ri, rd in enumerate(rows):
        rc = t.rows[ri + 1].cells
        bg = 'D6E4F0' if ri % 2 == 0 else 'FFFFFF'
        for ci, v in enumerate(rd):
            rc[ci].text = v
            shd(rc[ci], bg)
            if rc[ci].paragraphs[0].runs:
                rc[ci].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()
    return t


# ── Build Document ──────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ── TITLE PAGE ──────────────────────────────────────────────────────────────
for _ in range(3):
    doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('StoreWave')
tr.bold = True
tr.font.size = Pt(40)
tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run('Enterprise E-Commerce Platform')
sr.italic = True
sr.font.size = Pt(20)
sr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

sp2 = doc.add_paragraph()
sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr2 = sp2.add_run('Final System Documentation')
sr2.bold = True
sr2.font.size = Pt(16)
sr2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()

ip = doc.add_paragraph()
ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
ip.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}\n')
ip.add_run('Platform: ASP.NET Core 9.0 MVC\n')
ip.add_run('Live Demo: https://shopmartcommerce.runasp.net')

doc.add_page_break()

# ── TABLE OF CONTENTS ────────────────────────────────────────────────────────
ah(doc, 'Table of Contents', 1, '1F4E79')
TOC = [
    '1.  System Overview',
    '2.  Role-Based Features',
    '3.  System Architecture',
    '4.  Backend Architecture & Components',
    '5.  Data Flow in Backend',
    '6.  UML Class Diagram - Core Entities',
    '7.  Database Schema',
    '8.  Project Package Structure',
    '9.  Tools & Technologies',
    '10. Hosting & Deployment',
    '11. Getting Started (Local Development)',
    '12. Security & Authentication',
    '13. Real-Time Features (SignalR)',
    '14. Payment Integration (PayPal)',
    '15. Email & Notification System',
    '16. Team Members',
]
for item in TOC:
    doc.add_paragraph(item)
doc.add_page_break()

# ── 1. SYSTEM OVERVIEW ───────────────────────────────────────────────────────
ah(doc, '1. System Overview', 1, '1F4E79')
ap(doc,
   'StoreWave is a comprehensive, enterprise-grade e-commerce platform built using '
   'ASP.NET Core 9.0 MVC architecture. It provides a complete online shopping '
   'experience with a sophisticated multi-role management system, real-time '
   'capabilities, and a modern, responsive user interface.')
doc.add_paragraph()
at(doc, ['Feature', 'Description'], [
    ('Full Shopping Experience', 'Product browsing, search, filtering, cart management, and checkout'),
    ('PayPal Integration', 'Secure online payments via PayPal REST API'),
    ('Email Automation', 'OTP verification, order confirmations, notifications via Gmail SMTP'),
    ('Real-Time Updates', 'Live order tracking, cart sync, and order chat via SignalR WebSockets'),
    ('5-Role System', 'Admin, Customer, Supplier, Accountant, and Warehouse Manager'),
    ('Location Services', 'GPS-based address detection via Nominatim / OpenStreetMap'),
    ('Security', 'ASP.NET Identity with role-based authorization and anti-forgery tokens'),
    ('Analytics Dashboards', 'Revenue reports, sales analytics, and payment method statistics'),
])
doc.add_page_break()

# ── 2. ROLE-BASED FEATURES ───────────────────────────────────────────────────
ah(doc, '2. Role-Based Features', 1, '1F4E79')
ap(doc, 'StoreWave implements a Role-Based Access Control (RBAC) system with 5 '
        'distinct roles, each with dedicated dashboards and functionality.')
doc.add_paragraph()

ah(doc, '2.1  Admin - Full System Control', 2, '2E75B6')
at(doc, ['Feature', 'Description'], [
    ('Dashboard', 'Total revenue, orders, products, customers at a glance'),
    ('User Management', 'View all users, assign roles, activate/deactivate accounts'),
    ('Order Management', 'View all orders, update status: Pending > Processing > Shipped > Delivered'),
    ('Product Management', 'Full CRUD for all products across all suppliers'),
    ('Category Management', 'Create, edit, and delete product categories'),
    ('Order Chat', 'Real-time chat support with customers on specific orders'),
    ('Review Management', 'Monitor and manage product reviews'),
])

ah(doc, '2.2  Customer - Shopping Experience', 2, '2E75B6')
at(doc, ['Feature', 'Description'], [
    ('Home Page', 'Featured products, categories, and search functionality'),
    ('Product Browsing', 'Filter by category, search by name, view product details'),
    ('Shopping Cart', 'Add/remove items, update quantities, real-time cart sync'),
    ('Checkout', 'Cash on Delivery or PayPal payment options'),
    ('Order Tracking', 'View order history, track status, and order details'),
    ('Order Chat', 'Real-time chat with admin/support for specific orders'),
    ('Product Reviews', 'Rate and review purchased products'),
    ('Profile Management', 'Edit personal info, shipping address, and GPS location'),
])

ah(doc, '2.3  Supplier - Product Supply Management', 2, '2E75B6')
at(doc, ['Feature', 'Description'], [
    ('Dashboard', 'Total products, active listings, and out-of-stock alerts'),
    ('Product Management', 'Create, edit, delete own products with image upload'),
    ('Order Visibility', 'View orders containing their supplied products'),
    ('Image Management', 'Upload and manage product images'),
])

ah(doc, '2.4  Accountant - Financial Analytics', 2, '2E75B6')
at(doc, ['Feature', 'Description'], [
    ('Financial Dashboard', 'Total revenue, order count, average order value, today/monthly sales'),
    ('Sales Reports', 'Date-range filtered sales analysis'),
    ('Revenue Analytics', 'Monthly revenue trends over 12 months'),
    ('Payment Methods', 'Payment method distribution and statistics'),
    ('Order Details', 'Drill-down into individual order financials'),
])

ah(doc, '2.5  Warehouse Manager - Inventory Control', 2, '2E75B6')
at(doc, ['Feature', 'Description'], [
    ('Dashboard', 'Total products, low-stock alerts, out-of-stock count, total inventory'),
    ('Inventory Management', 'Full inventory listing sorted by stock level'),
    ('Stock Updates', 'Update stock quantities for individual products'),
    ('Low-Stock Alerts', 'Products with stock <= 10 are flagged for restocking'),
])
doc.add_page_break()

# ── 3. SYSTEM ARCHITECTURE ───────────────────────────────────────────────────
ah(doc, '3. System Architecture', 1, '1F4E79')
ap(doc, 'StoreWave follows a classic 3-Tier Architecture separating the Presentation, '
        'Business Logic, and Data layers, ensuring clean separation of concerns.')
doc.add_paragraph()
at(doc, ['Layer', 'Components', 'Responsibility'], [
    ('Presentation Layer', 'Razor Views (.cshtml), Bootstrap 5, jQuery, SignalR Client',
     'User interface, rendering, client-side interactivity'),
    ('Business Logic Layer', 'Controllers, Services, FluentValidation, AutoMapper, SignalR Hubs',
     'Business rules, validation, orchestration, real-time events'),
    ('Data Access Layer', 'Repository Pattern, Unit of Work, Entity Framework Core, SQL Server',
     'Data persistence, transactions, and query execution'),
    ('External Services', 'PayPal REST API, Gmail SMTP, Nominatim Geocoding',
     'Payment processing, email delivery, reverse geocoding'),
])
doc.add_page_break()

# ── 4. BACKEND ARCHITECTURE ──────────────────────────────────────────────────
ah(doc, '4. Backend Architecture & Components', 1, '1F4E79')
ap(doc, 'The backend follows Clean Architecture principles with clear separation of concerns.')
doc.add_paragraph()
at(doc, ['Component', 'Role', 'Pattern'], [
    ('Controllers (12)', 'Handle HTTP requests and route to services', 'MVC Controller'),
    ('Services (11)', 'Business logic, validation, orchestration', 'Service Layer'),
    ('Repositories (7)', 'Data access abstraction over EF Core', 'Repository Pattern'),
    ('Unit of Work', 'Transaction management across repositories', 'Unit of Work Pattern'),
    ('DTOs (7)', 'Data transfer between layers', 'Data Transfer Object'),
    ('ViewModels (12)', 'View-specific data models for Razor views', 'MVVM-inspired'),
    ('Validators (4)', 'Server-side input validation rules', 'FluentValidation'),
    ('Mappings', 'Entity <-> DTO conversion', 'AutoMapper Profile'),
    ('SignalR Hubs (4)', 'Real-time bidirectional communication', 'SignalR Hub'),
])

ah(doc, '4.1  Controllers', 2, '2E75B6')
at(doc, ['Controller', 'Responsibility'], [
    ('AccountController', 'Login, Register, OTP verification, password reset, access denied'),
    ('AdminController', 'Admin dashboard, user management, order management, order chat'),
    ('AccountantController', 'Financial reports, revenue analytics, payment statistics'),
    ('CartController', 'Shopping cart, checkout, order creation (Cash on Delivery)'),
    ('CategoriesController', 'Category CRUD operations'),
    ('HomeController', 'Home page and privacy page'),
    ('OrdersController', 'Customer order history and order details'),
    ('PaymentController', 'PayPal payment flow (create, capture, success, cancel)'),
    ('ProductsController', 'Product catalog browsing and CRUD'),
    ('ReviewsController', 'Product review submission and management'),
    ('SupplierController', 'Supplier dashboard and product management'),
    ('WarehouseController', 'Inventory listing and stock updates'),
])

ah(doc, '4.2  Services', 2, '2E75B6')
at(doc, ['Service', 'Key Responsibilities'], [
    ('ProductService', 'CRUD products, search, filter by category, image management'),
    ('OrderService', 'Create orders, update status, financial aggregations'),
    ('CartService', 'Add/remove/update cart items, get cart count, clear cart'),
    ('CustomerService', 'Profile management, account activation/deactivation'),
    ('CategoryService', 'Category CRUD, product-category association'),
    ('ReviewService', 'Review submission and rating calculation'),
    ('EmailService', 'Send transactional emails via Gmail SMTP'),
    ('EmailTemplateService', 'HTML email templates for OTP, order confirmation, etc.'),
    ('FileService', 'Upload and manage product/category images on the server'),
    ('PayPalService', 'Create and capture PayPal orders via REST API'),
    ('ChatService', 'Store and retrieve real-time order chat messages'),
])

ah(doc, '4.3  SignalR Hubs', 2, '2E75B6')
at(doc, ['Hub', 'Endpoint', 'Purpose'], [
    ('OrderHub', '/orderHub', 'Push real-time order status updates to customers and admins'),
    ('NotificationHub', '/notificationHub', 'Send live notifications (new orders, role events)'),
    ('CartHub', '/cartHub', 'Synchronize cart count badge in real-time'),
    ('ChatHub', '/chatHub', 'Enable two-way real-time chat between customer and admin per order'),
])
doc.add_page_break()

# ── 5. DATA FLOW ─────────────────────────────────────────────────────────────
ah(doc, '5. Data Flow in Backend', 1, '1F4E79')
ap(doc, 'The full request processing pipeline from the browser to the database and back:')
doc.add_paragraph()
at(doc, ['Step', 'Layer', 'Description'], [
    ('1', 'Browser', 'User sends HTTP request (page visit, form submit, AJAX call)'),
    ('2', 'Middleware', 'HTTPS redirect, static files, routing, session, authentication, authorization'),
    ('3', 'Controller', 'Matched controller action invoked; calls the appropriate service method'),
    ('4', 'Service', 'Business logic executed; input validated via FluentValidation'),
    ('5', 'Unit of Work', 'Database transaction opened; appropriate repository called'),
    ('6', 'Repository', 'LINQ queries composed over the EF Core DbContext'),
    ('7', 'EF Core', 'LINQ translated to SQL and executed against SQL Server'),
    ('8', 'SQL Server', 'Query runs; result set returned to EF Core as entity objects'),
    ('9', 'AutoMapper', 'Entities mapped to DTOs before returning to controller'),
    ('10', 'SignalR / Email', 'Side-effects: push WebSocket events or send email notifications'),
    ('11', 'Controller -> View', 'DTO/ViewModel passed to Razor view for server-side rendering'),
    ('12', 'Browser', 'HTML response (or JSON for AJAX) delivered to the user'),
])
doc.add_page_break()

# ── 6. UML CLASS DIAGRAM ─────────────────────────────────────────────────────
ah(doc, '6. UML Class Diagram - Core Entities', 1, '1F4E79')

ah(doc, '6.1  Customer  (extends IdentityUser<int>)', 2, '2E75B6')
at(doc, ['Property', 'Type', 'Description'], [
    ('Id', 'int', 'Primary key (inherited from IdentityUser)'),
    ('FirstName', 'string', 'Customer first name'),
    ('LastName', 'string', 'Customer last name'),
    ('Email', 'string', 'Unique email address (enforced by Identity)'),
    ('Address', 'string', 'Street address'),
    ('City', 'string', 'City of residence'),
    ('Country', 'string', 'Country'),
    ('PostalCode', 'string', 'Postal / ZIP code'),
    ('Latitude', 'double', 'GPS latitude for location services'),
    ('Longitude', 'double', 'GPS longitude for location services'),
    ('CreatedAt', 'DateTime', 'Account creation timestamp'),
    ('IsActive', 'bool', 'Whether the account is active'),
    ('FullName', 'string (computed)', 'FirstName + " " + LastName'),
    ('Orders', 'ICollection<Order>', 'Navigation: placed orders'),
    ('Reviews', 'ICollection<Review>', 'Navigation: written reviews'),
    ('CartItems', 'ICollection<CartItem>', 'Navigation: current shopping cart items'),
    ('SupplierProducts', 'ICollection<Product>', 'Navigation: products supplied (Supplier role)'),
])

ah(doc, '6.2  Product', 2, '2E75B6')
at(doc, ['Property', 'Type', 'Description'], [
    ('Id', 'int', 'Primary key'),
    ('Name', 'string', 'Product name'),
    ('Description', 'string', 'Product description'),
    ('Price', 'decimal', 'Regular price'),
    ('DiscountPrice', 'decimal', 'Discounted price (0 = no discount)'),
    ('StockQuantity', 'int', 'Available stock units'),
    ('ImageUrl', 'string', 'Path to product image file'),
    ('IsActive', 'bool', 'Whether the product is visible to customers'),
    ('IsFeatured', 'bool', 'Whether shown in home page featured section'),
    ('CreatedAt', 'DateTime', 'Product creation timestamp'),
    ('CategoryId', 'int (FK)', 'Foreign key to Category'),
    ('SupplierId', 'int (FK)', 'Foreign key to Customer (Supplier role)'),
    ('CurrentPrice', 'decimal (computed)', 'DiscountPrice if on sale, otherwise Price'),
    ('IsOnSale', 'bool (computed)', 'True when DiscountPrice > 0'),
    ('DiscountPercentage', 'int (computed)', 'Computed discount percentage'),
])

ah(doc, '6.3  Order', 2, '2E75B6')
at(doc, ['Property', 'Type', 'Description'], [
    ('Id', 'int', 'Primary key'),
    ('OrderNumber', 'string', 'Auto-generated unique order reference number'),
    ('OrderDate', 'DateTime', 'Date and time the order was placed'),
    ('SubTotal', 'decimal', 'Sum of all order item prices'),
    ('ShippingCost', 'decimal', 'Shipping fee applied to the order'),
    ('TotalAmount', 'decimal', 'SubTotal + ShippingCost'),
    ('Status', 'OrderStatus (enum)', 'Pending, Processing, Shipped, Delivered, or Cancelled'),
    ('ShippingAddress', 'string', 'Delivery address entered at checkout'),
    ('PaymentMethod', 'PaymentMethod (enum)', 'CashOnDelivery or PayPal'),
    ('Notes', 'string', 'Optional customer notes for the order'),
    ('CustomerId', 'int (FK)', 'Foreign key to the Customer who placed the order'),
    ('OrderItems', 'ICollection<OrderItem>', 'Navigation: line items in the order'),
    ('ChatMessages', 'ICollection<ChatMessage>', 'Navigation: support chat messages for this order'),
])

ah(doc, '6.4  Supporting Entities', 2, '2E75B6')
at(doc, ['Entity', 'Key Fields', 'Purpose'], [
    ('OrderItem', 'Id, OrderId, ProductId, ProductName, UnitPrice, Quantity, TotalPrice',
     'Individual product line item within an order'),
    ('CartItem', 'Id, CustomerId, ProductId, Quantity, AddedAt',
     'A product item added to a customer shopping cart'),
    ('Category', 'Id, Name, Description, ImageUrl, IsActive',
     'Product category for grouping and filtering products'),
    ('Review', 'Id, ProductId, CustomerId, Rating (1-5), Comment, CreatedAt',
     'Customer review and star rating for a product'),
    ('ChatMessage', 'Id, OrderId, SenderId, Message, SentAt, IsRead',
     'Real-time support chat message associated with an order'),
])

ah(doc, '6.5  Enumerations', 2, '2E75B6')
at(doc, ['Enum', 'Values', 'Used In'], [
    ('OrderStatus', 'Pending | Processing | Shipped | Delivered | Cancelled', 'Order.Status'),
    ('PaymentMethod', 'CashOnDelivery | PayPal', 'Order.PaymentMethod'),
])
doc.add_page_break()

# ── 7. DATABASE SCHEMA ───────────────────────────────────────────────────────
ah(doc, '7. Database Schema', 1, '1F4E79')
ap(doc, 'The database is managed by Entity Framework Core using Code-First migrations. '
        'The schema is automatically applied on application startup via context.Database.Migrate().')
doc.add_paragraph()
at(doc, ['Table', 'Primary Key', 'Foreign Keys', 'Description'], [
    ('AspNetUsers (Customer)', 'Id (int)', '--',
     'Stores all user accounts (Customer, Admin, Supplier, Accountant, Warehouse)'),
    ('AspNetRoles', 'Id (string)', '--', 'System roles: Admin, Customer, Supplier, Accountant, Warehouse'),
    ('AspNetUserRoles', 'UserId + RoleId', 'UserId, RoleId', 'Pivot table mapping users to roles'),
    ('Products', 'Id (int)', 'CategoryId, SupplierId', 'Product catalog with pricing and stock'),
    ('Categories', 'Id (int)', '--', 'Product category definitions'),
    ('Orders', 'Id (int)', 'CustomerId', 'Customer orders with status and payment info'),
    ('OrderItems', 'Id (int)', 'OrderId, ProductId', 'Individual product line items within orders'),
    ('CartItems', 'Id (int)', 'CustomerId, ProductId', 'Active shopping cart entries per customer'),
    ('Reviews', 'Id (int)', 'ProductId, CustomerId', 'Product reviews with star ratings'),
    ('ChatMessages', 'Id (int)', 'OrderId, SenderId', 'Order-specific support chat messages'),
])
doc.add_page_break()

# ── 8. PROJECT STRUCTURE ─────────────────────────────────────────────────────
ah(doc, '8. Project Package Structure', 1, '1F4E79')
struct = (
    'StoreWave/\n'
    '|\n'
    '+-- Controllers/              (12 MVC Controllers)\n'
    '|   +-- AccountController.cs       Authentication, OTP, Password Reset\n'
    '|   +-- AdminController.cs         Dashboard, Users, Orders, Chat\n'
    '|   +-- AccountantController.cs    Financial Reports, Revenue Analytics\n'
    '|   +-- CartController.cs          Cart, Checkout\n'
    '|   +-- CategoriesController.cs    Category CRUD\n'
    '|   +-- HomeController.cs          Home Page\n'
    '|   +-- OrdersController.cs        Customer Orders\n'
    '|   +-- PaymentController.cs       PayPal Payment Flow\n'
    '|   +-- ProductsController.cs      Product CRUD & Browsing\n'
    '|   +-- ReviewsController.cs       Product Reviews\n'
    '|   +-- SupplierController.cs      Supplier Product Management\n'
    '|   +-- WarehouseController.cs     Inventory & Stock Management\n'
    '|\n'
    '+-- Models/\n'
    '|   +-- Entities/  (8 domain entities)\n'
    '|   |   +-- Customer.cs    Product.cs    Order.cs    OrderItem.cs\n'
    '|   |   +-- CartItem.cs    Category.cs   Review.cs   ChatMessage.cs\n'
    '|   +-- Enums/\n'
    '|       +-- OrderStatus.cs    PaymentMethod.cs\n'
    '|\n'
    '+-- Views/                    (51 Razor Views)\n'
    '|   +-- Account/   (7 views)   Login, Register, Profile, OTP, Password\n'
    '|   +-- Admin/     (7 views)   Dashboard, Users, Orders, Chat\n'
    '|   +-- Accountant/(5 views)   Financial Reports, Revenue\n'
    '|   +-- Cart/      (3 views)   Cart, Checkout, Order Success\n'
    '|   +-- Categories/(5 views)   CRUD Views\n'
    '|   +-- Home/      (2 views)   Index, Privacy\n'
    '|   +-- Orders/    (2 views)   Order History, Details\n'
    '|   +-- Products/  (5 views)   Catalog, CRUD Views\n'
    '|   +-- Supplier/  (5 views)   Dashboard, Product Management\n'
    '|   +-- Warehouse/ (4 views)   Inventory, Stock Updates\n'
    '|   +-- Shared/    (4 views)   Layout, Login Partial, Error\n'
    '|\n'
    '+-- Services/\n'
    '|   +-- Interfaces/        (11 service contracts)\n'
    '|   +-- Implementations/  (11 service implementations)\n'
    '|\n'
    '+-- Repositories/\n'
    '|   +-- Interfaces/        (7 repository contracts)\n'
    '|   +-- Implementations/  (7 repository implementations)\n'
    '|\n'
    '+-- DTOs/           (7 Data Transfer Objects)\n'
    '+-- ViewModels/     (12 View Models)\n'
    '+-- Validators/     (4 FluentValidation rule sets)\n'
    '+-- Hubs/           (4 SignalR Hubs)\n'
    '+-- Data/           (ShopDbContext & DbSeeder)\n'
    '+-- Mappings/       (AutoMapper MappingProfile)\n'
    '+-- UnitOfWork/     (IUnitOfWork interface & implementation)\n'
    '+-- Migrations/     (EF Core migration files)\n'
    '+-- wwwroot/        (CSS, JavaScript, images, vendor libs)\n'
    '|\n'
    '+-- Program.cs         Entry point & Dependency Injection configuration\n'
    '+-- appsettings.json   Connection strings, SMTP, PayPal API keys\n'
    '+-- StoreWave.csproj   Project file with NuGet package references\n'
    '+-- StoreWave.sln      Visual Studio solution file\n'
)
cp = doc.add_paragraph()
cr = cp.add_run(struct)
cr.font.name = 'Courier New'
cr.font.size = Pt(9)
doc.add_page_break()

# ── 9. TOOLS & TECHNOLOGIES ──────────────────────────────────────────────────
ah(doc, '9. Tools & Technologies', 1, '1F4E79')

ah(doc, '9.1  Backend Stack', 2, '2E75B6')
at(doc, ['Technology', 'Version', 'Purpose'], [
    ('ASP.NET Core MVC', '9.0', 'Web framework & MVC pattern'),
    ('Entity Framework Core', '9.0', 'ORM & database migrations (Code-First)'),
    ('ASP.NET Identity', '9.0', 'Authentication & role-based authorization'),
    ('SignalR', '9.0', 'Real-time WebSocket communication'),
    ('FluentValidation', '11.3', 'Server-side input validation'),
    ('AutoMapper', '12.0', 'Object-to-object mapping (Entity <-> DTO)'),
    ('C#', '13', 'Primary programming language'),
])

ah(doc, '9.2  Frontend Stack', 2, '2E75B6')
at(doc, ['Technology', 'Version', 'Purpose'], [
    ('Razor Views (.cshtml)', '--', 'Server-side HTML templating'),
    ('Bootstrap', '5.3', 'Responsive UI framework'),
    ('jQuery', '3.7', 'DOM manipulation & AJAX requests'),
    ('JavaScript', 'ES6+', 'Client-side interactivity'),
    ('CSS3', '--', 'Custom styling & animations'),
    ('Font Awesome', '6.x', 'Icon library'),
    ('Leaflet.js', '--', 'Interactive maps for GPS location services'),
])

ah(doc, '9.3  Database & Storage', 2, '2E75B6')
at(doc, ['Technology', 'Purpose'], [
    ('SQL Server 2022', 'Primary relational database for all persistent data'),
    ('EF Core Migrations', 'Database schema versioning and automated migrations'),
    ('In-Memory Cache', 'Distributed caching layer (Redis-ready for production scale)'),
])

ah(doc, '9.4  External Services & APIs', 2, '2E75B6')
at(doc, ['Service', 'Purpose'], [
    ('PayPal REST API (v2)', 'Online payment processing: order creation & capture'),
    ('Gmail SMTP', 'Transactional email delivery (OTP, confirmations, notifications)'),
    ('Nominatim API', 'Reverse geocoding: convert GPS coordinates to address'),
    ('OpenStreetMap / Leaflet', 'Interactive location maps in the customer profile page'),
])

ah(doc, '9.5  Development Tools', 2, '2E75B6')
at(doc, ['Tool', 'Purpose'], [
    ('Visual Studio 2022', 'Primary IDE for development'),
    ('Git & GitHub', 'Version control and source code hosting'),
    ('SQL Server Management Studio (SSMS)', 'Database management and query execution'),
    ('NuGet', '.NET package management'),
])
doc.add_page_break()

# ── 10. HOSTING & DEPLOYMENT ─────────────────────────────────────────────────
ah(doc, '10. Hosting & Deployment', 1, '1F4E79')
ap(doc, 'StoreWave is deployed on MonsterASP.NET, a Windows-based ASP.NET hosting provider.')
doc.add_paragraph()
at(doc, ['Parameter', 'Value'], [
    ('Hosting Provider', 'MonsterASP.NET'),
    ('Runtime', 'ASP.NET Core 9.0'),
    ('Web Server', 'IIS (Internet Information Services)'),
    ('Database', 'SQL Server hosted on MonsterASP servers'),
    ('Domain / URL', 'https://shopmartcommerce.runasp.net'),
    ('Protocol', 'HTTPS with SSL certificate'),
    ('Deployment Method', 'Web Deploy from Visual Studio 2022'),
])
doc.add_paragraph()
ah(doc, '10.1  Deployment Steps', 2, '2E75B6')
for i, s in enumerate([
    'Build the project in Release mode: dotnet publish -c Release',
    'Configure the production connection string in appsettings.json',
    'Configure production SMTP and PayPal credentials in appsettings.json',
    'Publish via Web Deploy from Visual Studio 2022 to MonsterASP server',
    'On first startup, EF Core applies all pending migrations automatically',
    'DbSeeder auto-creates the system roles and default admin account',
], 1):
    doc.add_paragraph(f'{i}.  {s}')
doc.add_page_break()

# ── 11. GETTING STARTED ──────────────────────────────────────────────────────
ah(doc, '11. Getting Started (Local Development)', 1, '1F4E79')

ah(doc, '11.1  Prerequisites', 2, '2E75B6')
for x in [
    '.NET 9.0 SDK  -  https://dotnet.microsoft.com/download',
    'SQL Server 2019+ (LocalDB or Express edition)',
    'Visual Studio 2022 (recommended) or VS Code',
]:
    doc.add_paragraph(f'  -  {x}')

doc.add_paragraph()
ah(doc, '11.2  Setup Steps', 2, '2E75B6')
for i, s in enumerate([
    'Clone the repository from GitHub',
    'Open StoreWave.sln in Visual Studio 2022',
    'In appsettings.json, update DefaultConnection to point to your SQL Server instance',
    'Configure Gmail SMTP credentials (Email, Password) in appsettings.json',
    'Configure PayPal sandbox ClientId and Secret in appsettings.json',
    'Press F5 or run the application - EF Core will apply migrations and seed data automatically',
    'Navigate to https://localhost:{port} to access the application',
], 1):
    doc.add_paragraph(f'{i}.  {s}')

doc.add_paragraph()
ah(doc, '11.3  Default Admin Account (Auto-Seeded)', 2, '2E75B6')
at(doc, ['Role', 'Note'], [
    ('Admin', 'Created automatically by DbSeeder on first application run'),
])
doc.add_page_break()

# ── 12. SECURITY ─────────────────────────────────────────────────────────────
ah(doc, '12. Security & Authentication', 1, '1F4E79')
ap(doc, 'StoreWave uses ASP.NET Core Identity for authentication and role-based '
        'authorization, with the following security measures:')
doc.add_paragraph()
at(doc, ['Security Feature', 'Implementation Details'], [
    ('User Registration', 'Unique email enforced; minimum 6-character password requirement'),
    ('OTP Email Verification', 'Time-limited one-time passwords sent to user email for identity verification'),
    ('Password Reset Flow', 'Secure ASP.NET Identity token-based password reset via emailed link'),
    ('Role-Based Authorization', '[Authorize(Roles = "...")] attributes protect all role-restricted actions'),
    ('Cookie Authentication', 'Secure HttpOnly authentication cookies managed by ASP.NET Identity'),
    ('Anti-Forgery (CSRF)', 'CSRF protection via @Html.AntiForgeryToken() on all POST forms'),
    ('Account Activation', 'Admins can activate or deactivate any user account at any time'),
    ('Access Control Redirect', 'Unauthorized access attempts redirect to /Account/AccessDenied'),
])
doc.add_page_break()

# ── 13. SIGNALR ──────────────────────────────────────────────────────────────
ah(doc, '13. Real-Time Features (SignalR)', 1, '1F4E79')
ap(doc, 'ASP.NET Core SignalR provides persistent WebSocket connections enabling push '
        'communication from server to connected clients without polling.')
doc.add_paragraph()
at(doc, ['Hub', 'Client-Side Events', 'Triggered By'], [
    ('OrderHub', 'OrderStatusUpdated, NewOrderPlaced',
     'Admin changes order status; customer successfully places an order'),
    ('NotificationHub', 'NewNotification, RoleAssigned',
     'New order received; admin assigns a role to a user account'),
    ('CartHub', 'CartCountUpdated',
     'Item is added to or removed from the shopping cart'),
    ('ChatHub', 'ReceiveMessage, MessageSent',
     'Customer or admin sends a message in the order chat'),
])
doc.add_page_break()

# ── 14. PAYPAL ───────────────────────────────────────────────────────────────
ah(doc, '14. Payment Integration (PayPal)', 1, '1F4E79')
ap(doc, 'StoreWave integrates PayPal REST API (v2/checkout/orders) for secure online '
        'payment processing. The complete checkout flow is:')
doc.add_paragraph()
at(doc, ['Step', 'Description'], [
    ('1', 'Customer clicks "Pay with PayPal" on the checkout page'),
    ('2', 'POST /Payment/CreatePayPalOrder - PayPalService calls PayPal API to create an order and receives an approval URL'),
    ('3', 'Customer is redirected to PayPal secure payment page to review and approve'),
    ('4', 'After approval, PayPal redirects back to GET /Payment/PayPalSuccess with an order token'),
    ('5', 'PayPalService calls PayPal API to capture the payment; transaction is confirmed'),
    ('6', 'OrderService creates the order record in the database (PaymentMethod = PayPal)'),
    ('7', 'CartService clears all items from the customer cart'),
    ('8', 'EmailTemplateService generates and sends an order confirmation email to the customer'),
    ('9', 'NotificationHub pushes a "NewOrderPlaced" event to all connected admin clients'),
])
doc.add_paragraph()
ap(doc, 'If the customer cancels on PayPal\'s page, they are redirected to '
        'GET /Payment/PayPalCancel and no order is created in the database.', italic=True)
doc.add_page_break()

# ── 15. EMAIL ────────────────────────────────────────────────────────────────
ah(doc, '15. Email & Notification System', 1, '1F4E79')
ap(doc, 'StoreWave uses Gmail SMTP for all transactional email delivery. '
        'The EmailTemplateService generates rich HTML email bodies for each event type.')
doc.add_paragraph()
at(doc, ['Email Trigger', 'Recipient', 'Content Summary'], [
    ('Registration OTP', 'New customer', 'One-time password code for email address verification'),
    ('Password Reset', 'Existing customer', 'Secure ASP.NET Identity password-reset link'),
    ('Order Confirmation (CoD)', 'Customer', 'Full order summary: items, subtotal, shipping, total, estimated delivery'),
    ('Order Confirmation (PayPal)', 'Customer', 'Same as CoD confirmation with PayPal transaction reference'),
    ('Order Status Update', 'Customer', 'Notification of the new order status with tracking information'),
    ('Role Assignment', 'User', 'Notification that a new system role has been granted'),
])
doc.add_page_break()

# ── 16. TEAM MEMBERS ─────────────────────────────────────────────────────────
ah(doc, '16. Team Members', 1, '1F4E79')
ap(doc, 'The following table lists the project team members and their contributions:')
doc.add_paragraph()
at(doc, ['#', 'Name', 'Role / Responsibility'], [
    ('1', 'Ahmed Ramzy', 'Project Lead, Full-Stack Developer'),
    ('2', '', ''),
    ('3', '', ''),
    ('4', '', ''),
    ('5', '', ''),
])
doc.add_paragraph()
ap(doc, 'Note: Please fill in all team member names and their individual responsibilities above.',
   italic=True)

# ── SAVE ─────────────────────────────────────────────────────────────────────
OUT = (r'c:\Users\Ahmed Ramzy\OneDrive - Faculty of Computers & Artificial Intelligence'
       r'\Desktop\project MVC\StoreWave(Final Documentation).docx')
doc.save(OUT)
print('SUCCESS - Document saved to:')
print(OUT)
